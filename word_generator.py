# pyrefly: ignore [missing-import]
import io
import logging

logger = logging.getLogger(__name__)

class WordGenerator:
    """
    Generates Microsoft Word (.docx) documents from Legal Drafts.
    """
    
    @staticmethod
    def generate_draft_word(title: str, content: str, metadata: dict) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            document = Document()
            
            # Title
            heading = document.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            if metadata:
                for key, value in metadata.items():
                    p = document.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value))
            
            document.add_paragraph("_" * 50)
            
            # Content (rudimentary markdown parsing for now)
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    document.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    document.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    document.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    document.add_paragraph(line[2:], style='List Bullet')
                elif line.strip() == '':
                    continue
                else:
                    document.add_paragraph(line)
            
            # Save to BytesIO
            word_io = io.BytesIO()
            document.save(word_io)
            word_io.seek(0)
            return word_io.read()
            
        except ImportError as e:
            logger.error(f"python-docx is not installed: {e}")
            raise Exception("Word generation library is not installed.")
        except Exception as e:
            logger.error(f"Failed to generate Word document: {e}")
            raise e
